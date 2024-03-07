import docx
import os
from docx.shared import Pt 
from docx.enum.section import WD_SECTION
from template import delete_content, paste_content, change_content_style, get_content_style, check, replace_refrences
class Content_style:
    def __init__(self, font_size=12, font_name="Times New Roman"):
        self.font_size = font_size
        self.font_name = font_name

def main(journal_doc_path, template_doc_path):

    journal_doc = docx.Document(journal_doc_path)

    template_doc = docx.Document(template_doc_path)

    affiliation_end = ["email", "@", "corres", "introduction",'abstract']

    content_s = get_content_style(template_doc, "introduc")

    title = None
    authors = None
    affiliations = None
    email_and_corresponding_author = None
    abstract = None
    keywords = None

    #journal doc contents
    title, index = get_title(journal_doc)#0
    authors, index = get_authors(journal_doc, index) #1
    affiliations = get_affiliations(journal_doc, index, affiliation_end) #2
    email_and_corresponding_author = get_email_and_corresponding_author(journal_doc) #3
    abstract = get_abstract(journal_doc) #4 
    keywords = get_keywords(journal_doc) #5
    email = get_only_email(email_and_corresponding_author)
    corres_auth = get_only_corresponding_author(authors)

    #template doc contents
    title_t, title_t_index, next_index = get_title(template_doc, mode="dst")
    authors_t, authors_t_index, next_index = get_authors(template_doc, next_index, mode="dst")
    affiliations_t, affiliations_start_index, affiliation_end_index = get_affiliations(template_doc, next_index, affiliation_end, mode="dst")
    email_n_corr_t = get_email_and_corresponding_author(template_doc)
    email_t = get_only_email(email_n_corr_t)
    corr_author_t = get_only_corresponding_author(authors_t)
    abstract_t = get_abstract(template_doc)
    keywords_t = get_keywords(template_doc)

    #replacing contents
    replace_title(template_doc, title_t_index, title)
    replace_authors(template_doc, authors_t_index, authors, authors_t)
    replace_affiliations(template_doc, affiliations_start_index, affiliation_end_index, affiliations)
    replace_email(template_doc, email_t, email, corres_auth)
    replace_corresponding_author(template_doc, corr_author_t, corres_auth, email_t, email)
    replace_abstract(template_doc, abstract, abstract_t)
    replace_keyword(template_doc, keywords, keywords_t )

    print(content_s.font_name, content_s.font_size)
    
    for i in range(5):
        returns, content = check(template_doc, i)
        delete_content(template_doc, i)

    template_doc.add_paragraph("")

    for i in range(4):
        print(i)
        returns, content = check(template_doc, i)
        paste_content(template_doc, journal_doc, i)

    replace_refrences(journal_doc, template_doc)
    change_content_style(template_doc, content_s)

    journal_name = os.path.splitext(os.path.basename(journal_doc_path))[0]
    downloadable_filename = f"C.E - {journal_name}.docx"

    template_doc.save(f"output/{downloadable_filename}")
    return downloadable_filename

def get_title(doc,mode="src"):
    title = None
    for index, para in enumerate(doc.paragraphs):
        for run in para.runs:
            if run.bold or para.style.name.startswith('Heading') or  para.style.font.bold:

                title = para.text
                print("TITLE:",title, end='\n') # CHECKING
                if mode == "dst":
                    print(title, index, index + 1)
                    return title, index, index + 1 #the literal index and the next paragraph index
                print(title, index + 1)
                return title, index + 1


    print("TITLE not found") 
    return None

def replace_title(template_doc, title_t_index, source_title):
    if source_title == None:
        print("TITLE is None")
        return

    font_bold = None
    font_size = None
    font_name = None
    font_italic = None

    para = template_doc.paragraphs[title_t_index]
    for i, run in enumerate(para.runs):
        font_bold = run.font.bold
        font_size = run.font.size
        font_name = run.font.name
        font_italic = run.font.italic
        if i > 0:
            break

    para.text = source_title
    for i, run in enumerate(para.runs):
        run.font.bold = font_bold
        run.font.size = font_size
        run.font.name = font_name
        run.font.italic = font_italic

def get_authors(doc, start_index, mode="src"):
    authors = None
    for index, para in enumerate(doc.paragraphs, start=start_index):

        if not "corres" in doc.paragraphs[index].text.lower():
            if not doc.paragraphs[index].text == "":
                if doc.paragraphs[index].text.lower() == "abstract":
                    print("AUTHORS not found") 
                    return None, None
                authors = doc.paragraphs[index].text
                if mode == "dst":
                    return authors, index, index + 1
                return authors, index + 1
        else:
            break

    print("AUTHORS not found") 
    if mode == 'dst':
        return None, index, index
    return None, None

def replace_authors(template_doc, authors_t_index, source_author, template_author):
    if template_author == None:
        return
    if source_author == None:
        print("Source AUTHORS is None")
        return

    font_bold = None
    font_size = None
    font_name = None
    font_italic = None

    para = template_doc.paragraphs[authors_t_index]

    i = 0
    for run in para.runs:
        if run.text.strip():
            font_bold = run.font.bold
            font_size = run.font.size
            font_name = run.font.name
            font_italic = run.font.italic
            i += 1
        if i > 0:
            break

    para.text = source_author
    for i, run in enumerate(para.runs):
        run.font.bold = font_bold
        run.font.size = font_size
        run.font.name = font_name
        run.font.italic = font_italic

def get_affiliations(doc, start_index, affiliation_end, mode="src"):
    if start_index == None:
        print("START INDEX for AFFILIATIONS is None")
        return
    affs = []
    first_affiliation_index = 0
    last_affiliation_index = first_affiliation_index

    for index, para in enumerate(doc.paragraphs, start=start_index):
        if not doc.paragraphs[index].text == "":
            for end in affiliation_end:
                if end in doc.paragraphs[index].text.lower():
                    last_affiliation_index = index - 1
                    
                    if mode == "dst":
                        if len(affs) == 0:
                            return None, None, None

                        return affs, first_affiliation_index, last_affiliation_index
                    return affs
            
            if first_affiliation_index == 0:
                first_affiliation_index = index
            print(f"AFFILIATIONS {index}: {doc.paragraphs[index].text}", end='\n') # CHECKING
            affs.append(doc.paragraphs[index].text)

            if doc.paragraphs[index + 1].text == "":
                last_affiliation_index = index
                if mode == "dst":
                    return affs, first_affiliation_index, last_affiliation_index
                return affs


                if end in doc.paragraphs[index + 1].text.lower():
                    last_affiliation_index = index
                    if mode == "dst":
                        return affs, first_affiliation_index, last_affiliation_index
                    return affs


    print("AFFILIATIONS not found")   
    return None

def replace_affiliations(template_doc, affs_start_index, affs_end_index, source_affs):
    if affs_start_index == None:
        return
    
    print("affis", affs_start_index, affs_end_index)
    if source_affs == None:
        print("Source AFFILIATIONS is None")
        return
        
    current_affs_index = 0
    font_bold = None
    font_size = None
    font_name = None
    font_italic = None
    line_spacing = 1

    affs_new_line = [text + "\n" for text in source_affs]
    combined_affs = ''.join(affs_new_line)

    para = template_doc.paragraphs[affs_start_index]
    i = 0
    for run in para.runs:
        print(run.font.size)
        if run.text.strip():
            font_bold = run.font.bold
            font_size = run.font.size
            
            font_name = run.font.name
            font_italic = run.font.italic
            i += 1
        if i > 0:
            break
    line_spacing = template_doc.paragraphs[affs_start_index].paragraph_format.line_spacing

    if affs_start_index - affs_end_index == 0:
        para.text = combined_affs
        

        for run in para.runs:
            run.font.bold = font_bold
            run.font.size = font_size
            run.font.name = font_name
            run.font.italic = font_italic

        return
    else:
        para.text = combined_affs
        
        for run in para.runs:
            run.font.bold = font_bold
            run.font.size = font_size
            run.font.name = font_name
            run.font.italic = font_italic

        for i in range(affs_start_index + 1, affs_end_index + 1):
            para = template_doc.paragraphs[i]
            para.text = ""
        return

#     i = 0
#     for run in template_doc.paragraphs[affs_start_index].runs:
#         if run.text.strip():
#             font_bold = run.font.bold
#             font_size = run.font.size
#             font_name = run.font.name
#             font_italic = run.font.italic
#             i += 1
#         if i > 0:
#             break
#     line_spacing = template_doc.paragraphs[affs_start_index].paragraph_format.line_spacing

#     for j in range(affs_start_index, affs_end_index + 1):
#         para = template_doc.paragraphs[j]
#         amount = 0


#         if para.text.strip() and current_affs_index < len(source_affs) and not "corres" in  para.text.lower():
#             if affs_end_index - affs_start_index > 0:
#                 para = template_doc.paragraphs[affs_end_index - current_affs_index]
#             para.text = source_affs[len(source_affs) - 1 - current_affs_index]
#             for run in para.runs:
#                 run.font.bold = font_bold
#                 run.font.size = font_size
#                 run.font.name = font_name
#                 run.font.italic = font_italic

#             para.paragraph_format.line_spacing = line_spacing
#             amount += 1
#             current_affs_index += 1


#     while current_affs_index < len(source_affs):

#         new_para = template_doc.add_paragraph(source_affs[len(source_affs) - 1 - current_affs_index])
        
#         template_doc.paragraphs[affs_start_index].insert_paragraph_before(new_para.text)

#         new_para = template_doc.paragraphs[affs_start_index]

# # Set font properties for the runs in the new paragraph
#         for run in new_para.runs:
#             run.font.bold = font_bold
#             run.font.size = font_size
#             run.font.name = font_name
#             run.font.italic = font_italic
#         new_para.paragraph_format.line_spacing = line_spacing

#         current_affs_index += 1

def get_email_and_corresponding_author(doc,mode="src"):
    text = None
    for para in doc.paragraphs:
        if ("@" in para.text.lower()):
            # .append(
            # {f"name": "email", 
            # "content": para.text}
            # )
            text = para.text
            return text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "@" in para.text:
                        # .append(
                        #     {f"name": "email and author",
                        #     "content": para.text}
                        # )
                        text = para.text
                        if mode == "dst":
                            return text
                        return text
    # .append(
    # {f"name": "email and author", 
    # "content": "not found"}
    # )                       
    print("EMAIL and CORRESPONDING AUTHOR not found")
    return None

def get_only_email(text):
    if text == None:
        print("Cant get EMAIL because @ is not detected")
        return None
    at_index = text.find("@")
    
    if at_index != -1:

        start_index = text.rfind(" ", 0, at_index) + 1

        end_index = text.find(" ", at_index)
        if end_index == -1:
            end_index = len(text)

        substring = text[start_index:end_index]
        return substring
    else:
        print("There is no EMAIL")
        return None

def get_only_corresponding_author(authors):
    if authors == None:
        print("Cant get CORRESPONDING AUTHOR because AUTHORS is None")
        return
    author_list = authors.split(',')
    for name in author_list:
        if '*' in name:
            return name
    return author_list[0]

def replace_email(doc, template_email, source_email, src_author):
    if source_email == None:
        print("EMAIL is None")
        return
    if "*" in src_author:
        source_author = src_author[:-2]
    else:
        source_author = src_author[:-1]

    got_style = False
    font_bold = None
    font_size = None
    font_name = None
    font_italic = None

    for para in doc.paragraphs:
        if ("email" in para.text.lower() or "e-mail" in para.text.lower()) or "@" in para.text.lower():
            if not got_style:
                for i, run in enumerate(para.runs):
                    font_bold = run.font.bold
                    font_size = run.font.size
                    font_name = run.font.name
                    font_italic = run.font.italic
                    if i > 0:
                        got_style == True
                        break

            if source_email == None: # if there is no source email but template has source email and its the same para as corresponding author
                if "corresp" in para.text.lower():
                    para.text = f"Corresponding Author: {source_author}\nE-mail: -"
                    for run in para.runs:
                        run.font.bold = font_bold
                        run.font.size = font_size
                        run.font.name = font_name
                        run.font.italic = font_italic
                    return
                para.text = f"E-mail: -" # if it's different para than corresponding author
                for run in para.runs:
                    run.font.bold = font_bold
                    run.font.size = font_size
                    run.font.name = font_name
                    run.font.italic = font_italic
                return
            if "corresp" in para.text.lower(): # if email is the same line as corresponding author
                if source_email == None:
                    para.text = f"Corresponding Author: {source_author}\nE-mail: -"
                    for run in para.runs:
                        run.font.bold = font_bold
                        run.font.size = font_size
                        run.font.name = font_name
                        run.font.italic = font_italic
                    return
                para.text = f"Corresponding Author: {source_author}\nE-mail: {source_email}"
                for run in para.runs:
                    run.font.bold = font_bold
                    run.font.size = font_size
                    run.font.name = font_name
                    run.font.italic = font_italic
                return
            para.text = f"E-mail: {source_email}"
            for run in para.runs:
                run.font.bold = font_bold
                run.font.size = font_size
                run.font.name = font_name
                run.font.italic = font_italic
            return
        
    for para in doc.paragraphs:
        if "corresp" in para.text.lower():
            if not got_style:
                for i, run in enumerate(para.runs):
                    font_bold = run.font.bold
                    font_size = run.font.size
                    font_name = run.font.name
                    font_italic = run.font.italic
                    if i > 0:
                        got_style == True
                        break
            
            para.text = f"Corresponding Author: {source_author}\nE-mail: {source_email}"
            for run in para.runs:
                run.font.bold = font_bold
                run.font.size = font_size
                run.font.name = font_name
                run.font.italic = font_italic
            return

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                        words2 = para.text.split()
                        for j, word in enumerate(words2):
                            if template_email in word:
                                words2[j] = source_email
                                if not got_style:
                                    for i, run in enumerate(para.runs):
                                        font_bold = run.font.bold
                                        font_size = run.font.size
                                        font_name = run.font.name
                                        font_italic = run.font.italic
                                        if i > 0:
                                            got_style == True
                                            break
                                para.text = ''.join(words2)
                                for run in para.runs:
                                    run.font.bold = font_bold
                                    run.font.size = font_size
                                    run.font.name = font_name
                                    run.font.italic = font_italic
                                return
    
    print("EMAIL failed to be replaced")
    return

def replace_corresponding_author(template_doc, template_author, source_author, template_email, source_email):
    if template_author == None:
        return
    if source_author == None:
        print("Source COREESPONDING AUTHOR is None")
        return

    font_bold = None
    font_size = None
    font_name = None
    font_italic = None

    if "*" in template_author:
        new_t_auth = template_author[:-2]
    else:
        new_t_auth = template_author[:-1]

    if "*" in source_author:
        new_src_auth = source_author[:-2]
    else:
        new_src_auth = source_author[:-1]

    for para in template_doc.paragraphs:
        if not "introduction" in para.text.lower():
            if not "email" in para.text.lower() and not "e-mail" in para.text.lower():
                if "corresp" in para.text.lower():
                    for i, run in enumerate(para.runs):
                        font_bold = run.font.bold
                        font_size = run.font.size
                        font_name = run.font.name
                        font_italic = run.font.italic
                        if i > 0:
                            break
                    if source_author == None:
                        para.text = f"Corresponding Author: -"

                        for run in para.runs:
                            run.font.bold = font_bold
                            run.font.size = font_size
                            run.font.name = font_name
                            run.font.italic = font_italic
                        return
                    
                    if template_email == None:
                        para.text = f"Corresponding Author: {new_src_auth}\nE-mail: {source_email}"
                        for run in para.runs:
                            run.font.bold = font_bold
                            run.font.size = font_size
                            run.font.name = font_name
                            run.font.italic = font_italic
                        return
                    para.text = f"Corresponding Author: {new_src_auth}"
                    for run in para.runs:
                        run.font.bold = font_bold
                        run.font.size = font_size
                        run.font.name = font_name
                        run.font.italic = font_italic
                    return
        else:
            break
    
    for table in template_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if new_t_auth in para.text:
                        for i, run in enumerate(para.runs):
                            font_bold = run.font.bold
                            font_size = run.font.size
                            font_name = run.font.name
                            font_italic = run.font.italic
                            if i > 0:
                                break
                        
                        para.text = new_src_auth
                        for run in para.runs:
                            run.font.bold = font_bold
                            run.font.size = font_size
                            run.font.name = font_name
                            run.font.italic = font_italic
                        return
    print("failed replacing AUTHOR or AUTHOR already exist")
    return

def get_abstract(doc, mode="src"):
    abstract_text = []
    abstract = False
    amount = 0
    for index, para in enumerate(doc.paragraphs):

        if "keyword" in doc.paragraphs[index].text.lower() or "key word" in doc.paragraphs[index].text.lower():
            break

        if not abstract:
            if "abstract" in para.text.lower():
                abstract = True
                continue
        
        elif abstract and doc.paragraphs[index].text.strip():

            abstract_text.append(doc.paragraphs[index].text)


    if not abstract:
        abstract_cell_index = 0
        for table in doc.tables:
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if cell.text.lower() == 'abstract' and not abstract:
                        abstract = True
                        abstract_cell_index = j
                        break
                    elif abstract:

                        for p, cellpara in enumerate(row.cells[abstract_cell_index].paragraphs):
                            if "key word" in cellpara.text or "keyword" in cellpara.text  :
                                break

                            abstract_text.append(cellpara.text)
                        if p == len(row.cells[abstract_cell_index].paragraphs) - 1:
                            return abstract_text
                                
    if not abstract:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for e, para in enumerate(cell.paragraphs):
                        if not abstract:
                            if "abstract" in para.text.lower():
                                abstract = True
                                continue
                        else:
                            if "keyword" in para.text.lower() or "key word" in para.text.lower():
                                return abstract_text

                            abstract_text.append(para.text)

                            if e == len(cell.paragraphs) - 1:
                                return abstract_text
    if not abstract:
        print("ABSTRACT not found")                  
        return None
    else:
        return abstract_text

def replace_abstract(template_doc, src_abstract, template_abstract):
    if src_abstract == None:
        print("Source ABSTRACT is None")
        return

    font_bold = None
    font_size = None
    font_name = None
    font_italic = None

    found = False
    checking = True
    source_abstract = "".join(src_abstract)
    template_abstract_paras = 0
    for i, para in enumerate(template_doc.paragraphs):
        if "abstract" in para.text.lower() and not found:
            found = True
            continue
        
        if found:
            if para.text.strip() and len(template_abstract) == 1:
                para.text = source_abstract
                return
            elif para.text.strip() and len(template_abstract) > 1:
                if template_abstract_paras == 0:
                    for i, run in enumerate(para.runs):
                        font_bold = run.font.bold
                        font_size = run.font.size
                        font_name = run.font.name
                        font_italic = run.font.italic
                        if i > 0:
                            break

                    para.text = source_abstract
                    for run in para.runs:
                        run.font.bold = font_bold
                        run.font.size = font_size
                        run.font.name = font_name
                        run.font.italic = font_italic
                    template_abstract_paras += 1

                elif template_abstract_paras < len(template_abstract):
                    para.text = ""
                    template_abstract_paras += 1

    if not found:
        abstract_cell_index = 0
        for table in template_doc.tables:
            for o, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if cell.text.lower() == "abstract" and not found:
                        found = True
                        abstract_cell_index = j
                        break


                    if found:
                        for p, cellpara in enumerate(row.cells[abstract_cell_index].paragraphs):
                            if template_abstract_paras > len(template_abstract) - 1:
                                return
                            
                            if len(template_abstract) == 1:
                                if cellpara.text.strip():
                                    
                                    for i, run in enumerate(cellpara.runs):
                                        font_bold = run.font.bold
                                        font_size = run.font.size
                                        font_name = run.font.name
                                        font_italic = run.font.italic
                                        if i > 0:
                                            break

                                    cellpara.text = source_abstract

                                    for run in cellpara.runs:
                                        run.font.bold = font_bold
                                        run.font.size = font_size
                                        run.font.name = font_name
                                        run.font.italic = font_italic

                                    return
                            #maybe improve here
                            elif len(template_abstract) > 1:
                                if template_abstract_paras == 0 and cellpara.text.strip():
                                    for i, run in enumerate(cellpara.runs):
                                        font_bold = run.font.bold
                                        font_size = run.font.size
                                        font_name = run.font.name
                                        font_italic = run.font.italic
                                        if i > 0:
                                            break

                                    cellpara.text = source_abstract

                                    for run in cellpara.runs:
                                        run.font.bold = font_bold
                                        run.font.size = font_size
                                        run.font.name = font_name
                                        run.font.italic = font_italic

                                    template_abstract_paras += 1
                                elif template_abstract_paras < len(template_abstract):
                                    cellpara.text = ""
                                    template_abstract_paras += 1

    if not found:
        for table in template_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for i, para in enumerate(cell.paragraphs):
                        if "abstract" in para.text.lower() and not found:
                            found = True
                            continue

                        if found:
                            if para.text.strip() and len(template_abstract) == 1:
                                for i, run in enumerate(para.runs):
                                        font_bold = run.font.bold
                                        font_size = run.font.size
                                        font_name = run.font.name
                                        font_italic = run.font.italic
                                        if i > 0:
                                            break

                                para.text = source_abstract

                                for run in para.runs:
                                        run.font.bold = font_bold
                                        run.font.size = font_size
                                        run.font.name = font_name
                                        run.font.italic = font_italic
                                return
                            elif para.text.strip() and len(template_abstract) > 1:
                                if template_abstract_paras == 0:
                                    for i, run in enumerate(para.runs):
                                        font_bold = run.font.bold
                                        font_size = run.font.size
                                        font_name = run.font.name
                                        font_italic = run.font.italic
                                        if i > 0:
                                            break

                                    para.text = source_abstract

                                    for run in para.runs:
                                        run.font.bold = font_bold
                                        run.font.size = font_size
                                        run.font.name = font_name
                                        run.font.italic = font_italic

                                    template_abstract_paras += 1
                                elif template_abstract_paras < len(template_abstract):
                                    para.text = ""
                                    template_abstract_paras += 1

    if found:
        return
    else:
        print("failed replacing ABSTRACT")

def get_keywords(doc, mode="src"):
    keyword_text = []
    keyword = False
    amount = 0
    for index, para in enumerate(doc.paragraphs):
        if "keyword" in para.text.lower() or "key word" in para.text.lower():
            keyword = True

            keyword_text.append(para.text)
            return keyword_text
    
    for i,table in enumerate(doc.tables):
        for q,row in enumerate(table.rows):
            for w,cell in enumerate(row.cells):
                for e, para in enumerate(cell.paragraphs):
                    if "keyword" in para.text.lower() or "key word" in para.text.lower():
                        if len(cell.paragraphs) == 1:
                            #CHANGED
                            keyword_text.append(para.text)
                            return keyword_text
                        keyword = True   
                        continue
                    if keyword:
                        if para.text.strip():
                            amount += 1
                            keyword_text.append(para.text)
                        else:
                            if amount > 0:
                                return keyword_text

    print("KEYWORDS not found")
    return None

def replace_keyword(template_doc, source_keyword, template_keyword):
    if source_keyword == None:
        print("KEYWORD is None")
        return

    font_bold = None
    font_size = None
    font_name = None
    font_italic = None

    src_keywords = "".join(source_keyword)
    compare_keyword = template_keyword[0]

    if "keywords" in src_keywords.lower():
        keyword = src_keywords.replace("Keywords:", "")
    elif "key words" in src_keywords.lower():
        keyword = src_keywords.replace("Key words:", "")
    else:
        keyword = src_keywords

    for para in template_doc.paragraphs:
        if "keyword" in para.text.lower() or "key word" in para.text.lower():
            for i, run in enumerate(para.runs):
                font_bold = run.font.bold
                font_size = run.font.size
                font_name = run.font.name
                font_italic = run.font.italic
                if i > 0:
                    break
            para.text = f"Keywords: {keyword}"

            for run in para.runs:
                run.font.bold = font_bold
                run.font.size = font_size
                run.font.name = font_name
                run.font.italic = font_italic
            return

    found = False
    pasted = False
    t_keyword_index = 0
    
    for table in template_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "keyword" in para.text.lower() or "key word" in para.text.lower() and not found:
                        print("PARAGRAPH:", para.text)
                        if compare_keyword.lower() in para.text.lower():
                            print("KEYWORD FOUND")
                            for i, run in enumerate(para.runs):
                                font_bold = run.font.bold
                                font_size = run.font.size
                                font_name = run.font.name
                                font_italic = run.font.italic
                                if i > 0:
                                    break
                            if len(cell.paragraphs) == 1:
                                para.text = f"Keywords: {keyword}"
                                for run in para.runs:
                                    run.font.bold = font_bold
                                    run.font.size = font_size
                                    run.font.name = font_name
                                    run.font.italic = font_italic
                                return
                            para.text = f"Keywords:\n{keyword}"
                            for run in para.runs:
                                run.font.bold = font_bold
                                run.font.size = font_size
                                run.font.name = font_name
                                run.font.italic = font_italic
                            return
                        else:
                            print("COMPARE KEYWORD:",compare_keyword)
                            found = True
                            continue

                    elif found:
                        if not pasted:
                            if para.text.strip():
                                for i, run in enumerate(para.runs):
                                    font_bold = run.font.bold
                                    font_size = run.font.size
                                    font_name = run.font.name
                                    font_italic = run.font.italic
                                    if i > 0:
                                        break
                                para.text = keyword

                                for run in para.runs:
                                    run.font.bold = font_bold
                                    run.font.size = font_size
                                    run.font.name = font_name
                                    run.font.italic = font_italic

                                t_keyword_index += 1
                                pasted = True
                        else:
                            if t_keyword_index < len(template_keyword):
                                para.text = ''
                                t_keyword_index += 1
                            else:
                                # para.text = ''
                                return

    print('failed replacing KEYWORD')

