import docx
import os
from copy import deepcopy
from docx.oxml import parse_xml
import zipfile
from docx.shared import RGBColor
from docx.shared import Pt, Inches
import shutil
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement

# FUNCTIONS FOR REPLACING CONTENTS OF THE TEMPLATE

# identifying chapters and the format style in templates
class Content_style:
    def __init__(self, font_size=12, font_name="Times New Roman", line_spacing=1):
        self.font_size = font_size
        self.font_name = font_name
        self.line_spacing = line_spacing

def change_content_style(doc, style):
    found = False
    for i, para in enumerate(doc.paragraphs):
        words = para.text.split()
        word_amount = len(words)
        if not found:
            if word_amount > 30:
                found = True
        if found:
            for run in para.runs:
                run.font.size = Pt(style.font_size)
                run.font.name = style.font_name
            para.paragraph_format.line_spacing = style.line_spacing
        
def paste_content(target_doc, source_doc, chapter):
    chapter_order, _ = check(target_doc, chapter)
    # find the index where to start copying the content of a chapter 
    src_start, src_border = chapter_content_index(source_doc, chapter_order) 
    # find index of where the content should be pasted in the template file
    if chapter_order >= 5:
        dst_start, _ = chapter_content_index(target_doc, chapter_order, mode="paste") 
    else:
        dst_start, _ = chapter_content_index(target_doc, chapter_order)

    # find the index where to stop copying
    next_chapter_order = chapter_order + 1
    _, next_chapter_start = chapter_content_index(source_doc, next_chapter_order)

    insert_index = dst_start + 1  # Insert after the chapter title

    for i, element in enumerate(source_doc.element.body[src_border:next_chapter_start - 1]):
        target_doc.element.body.insert(insert_index + i, element)
    
def get_content_style(doc, chapter):
    
    found = False
    for para in doc.paragraphs:
        if not found:
            if chapter in para.text.lower():
                found = True
                continue
        else:
            for run in para.runs:
                if para.text.strip():
                    ls = para.paragraph_format.line_spacing
                    if run is not None:
                        if run.font.size is not None and run.font.name is not None:
                            return Content_style(font_size=run.font.size.pt, font_name=run.font.name, line_spacing=ls)
                        else:
                            return Content_style()
    return Content_style()

# check if a paragraph is a list
def is_numbered(paragraph):
    return paragraph.style.name.startswith("List") and paragraph.style.name != "List Paragraph"

# MAKE IT LISTED WITH NUMBER
def replace_refrences(journal_doc, template_doc):
    _,_, para_index_j = chapter_content_index(journal_doc, 4, mode="paste")
    _,_, para_index_t = chapter_content_index(template_doc, 4, mode="paste")
    refs = []
    found = False
    for i in range(para_index_j - 1, len(journal_doc.paragraphs)):
        para = journal_doc.paragraphs[i]
        refs.append(para.text)
    list_number_style = template_doc.styles.add_style('List Number', WD_STYLE_TYPE.PARAGRAPH)
    list_number_style.paragraph_format.left_indent = Inches(0.5)  # Adjust as needed
    list_number_style.paragraph_format.first_line_indent = Inches(-0.25)  # Adjust as needed
    list_number_style.base_style = template_doc.styles['Normal']
    for text in refs:
        template_doc.add_paragraph(text, style='List Number')
        # p.style = 'List Paragraph'


# function to find the index of the title of a chapter and the start of its content
def chapter_content_index(doc, chapter, mode="delete"):
    chapter_order, content = check(doc, chapter)
    if chapter_order > 4:
        return len(doc.element.body) - 1, len(doc.element.body) - 1
        
    para_index = 0 # tracking paragraphs
    title_index = 0 # to store the index of the title
    table_index =  0 # tracking tables
    found = False 

    for i, element in enumerate(doc.element.body):
        if element.tag.endswith('p'):
            para_index += 1

            para = doc.paragraphs[para_index - 1]

            if not found:
                # Check if the paragraph contains the content
                if content.lower() in para.text.lower():
                    # Check if the paragraph is numbered
                    if is_numbered(para):
                        # Extract text excluding numbering
                        para_text = para.text.split(' ', 1)[1].strip()
                        if para_text.lower() == content.lower():
                            # Chapter title found in a numbered paragraph
                            found = True
                            title_index = i
                            break
                    else:
                        for j, run in enumerate(para.runs):
                            
                            if j < 3:
                                # Check if the paragraph is styled as a heading or if it's bold
                                if run.bold or para.style.name.startswith('Heading') or para.style.font.bold:
                                    # extra checking to reduce the chance of mistakes
                                    for substring in ["ntro", "etho", "efer", "esul", "onclu"]: # content was here
                                        if substring in run.text.lower():
                                            if substring == "esul":
                                                if "test" in run.text.lower():
                                                    continue
                                            found = True
                                            title_index = i
                                            break
                                    if found:
                                        break
                            else:
                                break
            else:
                if mode == "delete":
                    return title_index, i
                if mode == "paste":
                    if content == "references":
                        return title_index, title_index + 1, para_index
                    return title_index, title_index + 1
        # checking if there is any title in a table
        elif element.tag.endswith('tbl'):
            table_index += 1
            if content in doc.tables[table_index - 1].rows[0].cells[0].paragraphs[0].text.lower() and doc.tables[table_index - 1].rows[0].cells[0].paragraphs[0].runs[0].bold and not found:
                title_index = i
                found = True
                return title_index, i + 1
            elif found:
                return title_index, i
            

def delete_content(doc, chapter):
    chapter_order, content = check(doc, chapter)
    title_index, content_i = chapter_content_index(doc, chapter_order)
    border_index, content_border_i = chapter_content_index(doc, (chapter_order + 1))
    for i in range(border_index - 1, content_i -1, -1):
        doc.element.body.remove(doc.element.body[i])
    
def check(doc, chapter):
    for i in range(chapter, len(chapter_names)):
        if chapter == 0:
            return 0, "introduction"
        for para in doc.paragraphs:
            if chapter_names[i] in para.text.lower():
                for j, run in enumerate(para.runs):
                        if run.bold or para.style.name.startswith('Heading'):
                                    return i , chapter_names[i]
        for table in doc.tables:
            if chapter_names[i] in table.rows[0].cells[0].paragraphs[0].text.lower() and table.rows[0].cells[0].paragraphs[0].runs[0].bold:
                return i, chapter_names[i]
    return len(chapter_names), "end"

chapter_names = ["introduction" , "method", "result", "conclusion", "references"]

